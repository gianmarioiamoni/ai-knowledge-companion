#!/usr/bin/env python3
"""
Script per creare file DOC e DOCX di test
Richiede: pip install python-docx
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx non installato")
    print("Installa con: pip install python-docx")
    exit(1)

def create_test_docx():
    """Crea un file DOCX di test"""
    doc = Document()
    
    # Titolo
    title = doc.add_heading('AI Knowledge Companion - Test Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sottotitolo
    subtitle = doc.add_paragraph('This is a test document for the AI Knowledge Companion application.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.italic = True
    
    # Introduzione
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph(
        'The AI Knowledge Companion is a sophisticated platform that combines artificial '
        'intelligence with knowledge management. This document serves as a test file for '
        'validating the DOCX parsing capabilities using LangChain\'s DocxLoader.'
    )
    
    # Features
    doc.add_heading('Features', level=1)
    features = [
        'Document Upload: Support for multiple file formats including PDF, TXT, MD, DOC, and DOCX',
        'RAG System: Retrieval-Augmented Generation for intelligent responses',
        'Vector Search: Using pgvector for semantic similarity search',
        'OpenAI Integration: Text embeddings and GPT-4o-mini for conversations',
        'Tutor Management: Create and manage AI tutors with custom knowledge bases'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Technical Details
    doc.add_heading('Technical Details', level=1)
    doc.add_paragraph('The system uses several key technologies:')
    technologies = [
        'Next.js 15 for the frontend framework',
        'Supabase for backend and database',
        'LangChain for document processing',
        'OpenAI for embeddings and chat completions',
        'TypeScript for type safety'
    ]
    for i, tech in enumerate(technologies, 1):
        doc.add_paragraph(f'{i}. {tech}', style='List Number')
    
    # Document Processing Pipeline
    doc.add_heading('Document Processing Pipeline', level=1)
    doc.add_paragraph('When a document is uploaded, it goes through the following steps:')
    steps = [
        'File validation and type detection',
        'Text extraction using appropriate parser (WebPDFLoader for PDF, DocxLoader for Word documents)',
        'Text chunking with RecursiveCharacterTextSplitter (1000 characters, 120 overlap)',
        'Embedding generation using OpenAI text-embedding-3-small model',
        'Storage in Supabase with pgvector for similarity search'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion = doc.add_paragraph(
        'This test document contains enough varied content to properly test the DOCX parsing, '
        'chunking, and embedding generation process. The content includes headers, lists, '
        'technical terms, and multiple paragraphs to ensure comprehensive testing.'
    )
    
    # Footer
    footer = doc.add_paragraph('\nEnd of Test Document')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer.runs[0]
    footer_format.bold = True
    footer_format.font.color.rgb = RGBColor(128, 128, 128)
    
    # Salva il documento
    doc.save('test-document.docx')
    print("✓ Created: test-document.docx")

def create_simple_docx():
    """Crea un file DOCX semplice per test rapidi"""
    doc = Document()
    
    doc.add_heading('Simple Test Document', 0)
    doc.add_paragraph('This is a simple test document with minimal content.')
    doc.add_paragraph('It contains only a few paragraphs to test basic DOCX parsing.')
    doc.add_paragraph('The AI Knowledge Companion should be able to extract this text correctly.')
    
    doc.save('simple-test.docx')
    print("✓ Created: simple-test.docx")

if __name__ == '__main__':
    try:
        create_test_docx()
        create_simple_docx()
        print("\n✅ All test documents created successfully!")
        print("\nNote: DOC format requires Microsoft Word or LibreOffice to convert.")
        print("You can open the DOCX files and save them as .doc format if needed.")
    except Exception as e:
        print(f"❌ Error creating documents: {e}")

